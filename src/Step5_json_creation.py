"""Step 5: Batch JSON creation from drug reports via Responses API + Instructor.

This script reads:
- a prompt template (.docx),
- a directory of drug report markdown files (*.md),
- a directory of per-drug pathway mapping JSON files (<drug>.json),

and produces one validated structured JSON output per drug:
- <output_dir>/<drug>.json

Notes:
- Drug name is inferred from markdown filename stem.
- Each drug expects pathway JSON at <pathway_dir>/<drug>.json.
- Missing pathway files are skipped with warning.
- Responses API is used for generation.
- Instructor + Pydantic enforce output schema and retries.
"""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
from time import sleep
from typing import Any, Dict, List, Optional

from openai import OpenAI
from pydantic import BaseModel, ConfigDict, Field

# Optional .env loading
try:
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:
    pass

# .docx reader
try:
    from docx import Document
except ImportError as exc:
    raise ImportError("python-docx is required. Install with: pip install python-docx") from exc

# Instructor is optional here; some versions do not expose `.responses`.
try:
    import instructor
except ImportError as exc:
    raise ImportError(
        "instructor is required. Install with: pip install instructor"
    ) from exc


DEFAULT_MODEL = "o4-mini-deep-research"
DEFAULT_TEMPERATURE = 0.2
DEFAULT_MAX_OUTPUT_TOKENS = 16000
DEFAULT_FILE_PATTERN = "*.md"
DEFAULT_MAX_RETRIES = 3
THROTTLE_SECONDS = 0.3


class StructuredDrugOutput(BaseModel):
    """Prompt3a-aligned canonical top-level schema with strict key control."""

    model_config = ConfigDict(extra="forbid")

    drug_name: str
    cancer_indication: Optional[str] = None
    drug_category: Optional[str] = None
    drug_class: Optional[str] = None
    moa: Optional[str] = None
    chembl_id: Optional[str] = None
    drugbank_id: Optional[str] = None
    synonyms: List[str] = Field(default_factory=list)
    primary_targets: List[str] = Field(default_factory=list)
    pathway_sets: List[str] = Field(default_factory=list)
    pathway_sets_annotations: Dict[str, Dict[str, Any]] = Field(default_factory=dict)
    sensitivity_genes_up: List[str] = Field(default_factory=list)
    sensitivity_genes_down: List[str] = Field(default_factory=list)
    resistance_genes_up: List[str] = Field(default_factory=list)
    resistance_genes_down: List[str] = Field(default_factory=list)
    sensitivity_genes_up_annotations: List[Dict[str, Any]] = Field(default_factory=list)
    sensitivity_genes_down_annotations: List[Dict[str, Any]] = Field(default_factory=list)
    resistance_genes_up_annotations: List[Dict[str, Any]] = Field(default_factory=list)
    resistance_genes_down_annotations: List[Dict[str, Any]] = Field(default_factory=list)
    kg_gene_relationships: List[Dict[str, Any]] = Field(default_factory=list)
    contraindications: List[str] = Field(default_factory=list)
    citations: List[Any] = Field(default_factory=list)
    notes: str = ""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Step 5 batch JSON creation using Responses API + Instructor. "
            "Deep-research options: o4-mini-deep-research (default), o3-deep-research."
        )
    )
    parser.add_argument("--prompt-docx", required=True, help="Path to prompts/prompt3a.docx")
    parser.add_argument("--md-dir", required=True, help="Directory containing per-drug markdown files")
    parser.add_argument(
        "--pathway-dir",
        required=True,
        help="Directory containing per-drug pathway JSON files named <drug>.json",
    )
    parser.add_argument("--output-dir", required=True, help="Directory to write <drug>.json outputs")

    parser.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help=(
            "OpenAI model for Responses API "
            "(recommended: o4-mini-deep-research or o3-deep-research)."
        ),
    )
    parser.add_argument("--temperature", type=float, default=DEFAULT_TEMPERATURE)
    parser.add_argument("--max-output-tokens", type=int, default=DEFAULT_MAX_OUTPUT_TOKENS)
    parser.add_argument("--file-pattern", default=DEFAULT_FILE_PATTERN, help="Markdown glob pattern")
    parser.add_argument("--max-retries", type=int, default=DEFAULT_MAX_RETRIES)
    return parser.parse_args()


def validate_inputs(args: argparse.Namespace) -> None:
    prompt_path = Path(args.prompt_docx)
    md_dir = Path(args.md_dir)
    pathway_dir = Path(args.pathway_dir)
    output_dir = Path(args.output_dir)

    if not prompt_path.exists() or not prompt_path.is_file():
        raise FileNotFoundError(f"Prompt docx not found: {prompt_path}")
    if not md_dir.exists() or not md_dir.is_dir():
        raise FileNotFoundError(f"Markdown directory not found: {md_dir}")
    if not pathway_dir.exists() or not pathway_dir.is_dir():
        raise FileNotFoundError(f"Pathway directory not found: {pathway_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)


def init_clients() -> tuple[OpenAI, Any]:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY not found. Set it with environment variable or .env file."
        )

    openai_client = OpenAI(api_key=api_key)

    mode = None
    for name in ("RESPONSES_TOOLS", "RESPONSES_JSON", "JSON"):
        if hasattr(instructor.Mode, name):
            mode = getattr(instructor.Mode, name)
            break

    instructor_client = (
        instructor.from_openai(openai_client, mode=mode)
        if mode is not None
        else instructor.from_openai(openai_client)
    )
    return openai_client, instructor_client


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def save_json(obj: Any, path: Path) -> None:
    with path.open("w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    full_text: List[str] = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)


def read_markdown(path: Path) -> str:
    with path.open("r", encoding="utf-8") as f:
        return f.read()


def extract_pathway_names(pathway_json_path: Path) -> List[str]:
    pathway_data = load_json(pathway_json_path)
    pathway_names: List[str] = []
    for _, row_data in pathway_data.items():
        if isinstance(row_data, dict) and "Mapped MSigDB Pathway Name" in row_data:
            pathway_name = row_data["Mapped MSigDB Pathway Name"]
            if isinstance(pathway_name, str):
                cleaned = pathway_name.strip()
                if cleaned and cleaned not in pathway_names:
                    pathway_names.append(cleaned)
    return pathway_names


def build_user_prompt(
    prompt_template: str,
    drug_report: str,
    drug_name: str,
    pathway_list: List[str],
    correction_note: Optional[str] = None,
) -> str:
    pathway_list_str = "\n".join([f"  {i + 1}. {name}" for i, name in enumerate(pathway_list)])

    prompt = f"""{prompt_template}

---

DRUG NAME: {drug_name}

<PATHWAY_LIST>
{pathway_list_str}
</PATHWAY_LIST>

DRUG REPORT:
{drug_report}

---

Based on the prompt instructions above, the pathway list, and the drug report provided, generate a comprehensive structured JSON output.
The JSON should capture all relevant pathway-drug interactions, mechanisms, clinical evidence, and classifications from the report.
Use the pathway names from <PATHWAY_LIST> as the canonical pathway identifiers.

Return ONLY valid JSON (no markdown code blocks, no explanation outside JSON).
"""
    if correction_note:
        prompt += f"\n\nVALIDATION FEEDBACK (fix and regenerate):\n{correction_note}\n"
    return prompt


def _normalize_instructor_result(result: Any) -> Dict[str, Any]:
    if isinstance(result, StructuredDrugOutput):
        return result.model_dump()
    if hasattr(result, "model_dump"):
        data = result.model_dump()
        return StructuredDrugOutput.model_validate(data).model_dump()
    if isinstance(result, dict):
        return StructuredDrugOutput.model_validate(result).model_dump()
    raise TypeError(f"Unexpected instructor result type: {type(result)}")


def _extract_response_text(response: Any) -> str:
    """Extract text from Responses API result across SDK variants."""
    text = getattr(response, "output_text", None)
    if isinstance(text, str) and text.strip():
        return text

    output = getattr(response, "output", None)
    if isinstance(output, list):
        chunks: List[str] = []
        for item in output:
            content = getattr(item, "content", None)
            if not isinstance(content, list):
                continue
            for part in content:
                part_text = getattr(part, "text", None)
                if isinstance(part_text, str):
                    chunks.append(part_text)
        if chunks:
            return "\n".join(chunks)
    return ""


def _is_deep_research_model(model: str) -> bool:
    return "deep-research" in (model or "").lower()


def call_responses_with_retry(
    openai_client: OpenAI,
    instructor_client: Any,
    model: str,
    temperature: float,
    max_output_tokens: int,
    prompt_template: str,
    drug_report: str,
    drug_name: str,
    pathway_list: List[str],
    max_retries: int,
) -> Dict[str, Any]:
    system_prompt = (
        "You are an expert in cancer biology, pharmacology, and pathway analysis. "
        "Your task is to extract and structure information from drug reports into comprehensive JSON format. "
        "Be precise, scientifically accurate, and ensure the JSON exactly matches the requested schema. "
        "Return only valid JSON."
    )

    correction_note: Optional[str] = None
    last_error: Optional[Exception] = None
    use_server_json_schema = not _is_deep_research_model(model)

    for attempt in range(max_retries):
        user_prompt = build_user_prompt(
            prompt_template=prompt_template,
            drug_report=drug_report,
            drug_name=drug_name,
            pathway_list=pathway_list,
            correction_note=correction_note,
        )

        try:
            # Primary path: Responses API with strict JSON schema.
            request_kwargs: Dict[str, Any] = {
                "model": model,
                "input": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": temperature,
                "max_output_tokens": max_output_tokens,
            }

            # Deep-research models require at least one tool.
            if _is_deep_research_model(model):
                request_kwargs["tools"] = [{"type": "web_search_preview"}]

            if use_server_json_schema:
                request_kwargs["text"] = {
                    "format": {
                        "type": "json_schema",
                        "name": "structured_drug_output",
                        "schema": StructuredDrugOutput.model_json_schema(),
                        # Keep server-side schema guidance non-strict because this
                        # payload includes free-form nested dict sections that do not
                        # satisfy OpenAI strict-schema additionalProperties=false rules.
                        # Final strictness is enforced by local Pydantic validation below.
                        "strict": False,
                    }
                }

            result = openai_client.responses.create(**request_kwargs)
            response_text = _extract_response_text(result)
            if not response_text.strip():
                raise ValueError("Responses API returned empty text output.")
            validated = StructuredDrugOutput.model_validate_json(response_text).model_dump()
            validated["drug_name"] = drug_name
            return validated
        except Exception as exc:  # API + validation failures are both retried.
            last_error = exc
            err_text = str(exc)
            print(f"OpenAI/instructor error (attempt {attempt + 1}/{max_retries}): {exc}")

            # Deep-research models can reject text.format=json_schema.
            # Disable server-side schema formatting and continue retrying with local validation.
            if (
                use_server_json_schema
                and "text.format" in err_text
                and "not supported" in err_text
            ):
                use_server_json_schema = False
                correction_note = (
                    "The model does not support server-side json_schema response formatting. "
                    "Return raw valid JSON that matches the required schema exactly."
                )
                if attempt < max_retries - 1:
                    sleep(2**attempt)
                continue

            correction_note = (
                "The previous response failed schema validation or JSON parsing. "
                f"Error: {str(exc)}. "
                "Regenerate and strictly follow the required JSON schema with all expected top-level keys."
            )
            if attempt < max_retries - 1:
                sleep(2**attempt)

    raise RuntimeError(f"Failed to generate valid JSON for {drug_name}: {last_error}")


def process_single_drug(
    openai_client: OpenAI,
    instructor_client: Any,
    args: argparse.Namespace,
    prompt_template: str,
    md_path: Path,
    pathway_dir: Path,
    output_dir: Path,
) -> str:
    drug_name = md_path.stem
    pathway_json_path = pathway_dir / f"{drug_name}.json"

    if not pathway_json_path.exists():
        return "skipped"

    drug_report = read_markdown(md_path)
    pathway_list = extract_pathway_names(pathway_json_path)

    if not pathway_list:
        print(f"WARNING: No pathways extracted for '{drug_name}'. Proceeding with empty list.")

    result = call_responses_with_retry(
        openai_client=openai_client,
        instructor_client=instructor_client,
        model=args.model,
        temperature=args.temperature,
        max_output_tokens=args.max_output_tokens,
        prompt_template=prompt_template,
        drug_report=drug_report,
        drug_name=drug_name,
        pathway_list=pathway_list,
        max_retries=args.max_retries,
    )

    output_path = output_dir / f"{drug_name}.json"
    save_json(result, output_path)
    print(f"Saved: {output_path}")
    return "processed"


def run_batch_pipeline(args: argparse.Namespace) -> None:
    validate_inputs(args)

    prompt_path = Path(args.prompt_docx)
    md_dir = Path(args.md_dir)
    pathway_dir = Path(args.pathway_dir)
    output_dir = Path(args.output_dir)

    openai_client, instructor_client = init_clients()

    prompt_template = read_docx(prompt_path)
    md_files = sorted(md_dir.glob(args.file_pattern))

    print("=" * 70)
    print("Step 5: Batch JSON Creation")
    print("=" * 70)
    print(f"Model: {args.model}")
    print(f"Prompt: {prompt_path}")
    print(f"Markdown dir: {md_dir}")
    print(f"Pathway dir: {pathway_dir}")
    print(f"Output dir: {output_dir}")
    print(f"File pattern: {args.file_pattern}")
    print(f"Max retries: {args.max_retries}")
    print(f"Throttle: {THROTTLE_SECONDS}s per drug")
    print("-" * 70)

    if not md_files:
        print(f"No markdown files found in {md_dir} with pattern '{args.file_pattern}'.")
        return

    processed = 0
    skipped = 0
    failed = 0

    for index, md_path in enumerate(md_files, start=1):
        drug_name = md_path.stem
        pathway_json_path = pathway_dir / f"{drug_name}.json"

        print(f"\n[{index}/{len(md_files)}] Drug: {drug_name}")
        if not pathway_json_path.exists():
            print(f"WARNING: Missing pathway file, skipping: {pathway_json_path}")
            skipped += 1
            sleep(THROTTLE_SECONDS)
            continue

        try:
            status = process_single_drug(
                openai_client=openai_client,
                instructor_client=instructor_client,
                args=args,
                prompt_template=prompt_template,
                md_path=md_path,
                pathway_dir=pathway_dir,
                output_dir=output_dir,
            )
            if status == "processed":
                processed += 1
            else:
                skipped += 1
        except Exception as exc:
            failed += 1
            print(f"ERROR: Failed for '{drug_name}': {exc}")
        finally:
            sleep(THROTTLE_SECONDS)

    print("\n" + "=" * 70)
    print("Step 5 Summary")
    print("=" * 70)
    print(f"Total markdown files discovered: {len(md_files)}")
    print(f"Processed: {processed}")
    print(f"Skipped: {skipped}")
    print(f"Failed: {failed}")


def main() -> None:
    args = parse_args()
    run_batch_pipeline(args)


if __name__ == "__main__":
    main()
